"""
server.py — Market Research веб-сервер
pip install flask httpx reportlab python-docx matplotlib
python server.py  →  http://localhost:5000
"""

import os, asyncio, io, re, json, base64, urllib.request
from datetime import datetime
import httpx
from flask import Flask, request, jsonify, send_file

app = Flask(__name__)

PERPLEXITY_API_KEY = os.environ.get("PERPLEXITY_API_KEY", "")
PERPLEXITY_URL     = "https://api.perplexity.ai/chat/completions"
AUTHOR             = "Создано Treyner Igor, e-mail: i.treyner@yandex.ru"

# ---------------------------------------------------------------------------
# DejaVu fonts for PDF Cyrillic support
# ---------------------------------------------------------------------------
FONT_DIR  = "/tmp/mr_fonts"
FONT_URLS = {
    "DejaVuSans":      "https://github.com/dejavu-fonts/dejavu-fonts/raw/master/ttf/DejaVuSans.ttf",
    "DejaVuSans-Bold": "https://github.com/dejavu-fonts/dejavu-fonts/raw/master/ttf/DejaVuSans-Bold.ttf",
}

def ensure_fonts():
    os.makedirs(FONT_DIR, exist_ok=True)
    for name, url in FONT_URLS.items():
        path = f"{FONT_DIR}/{name}.ttf"
        if not os.path.exists(path):
            try:
                urllib.request.urlretrieve(url, path)
            except Exception as e:
                print(f"Font download failed ({name}): {e}")

_fonts_registered = False

def register_pdf_fonts():
    global _fonts_registered
    if _fonts_registered:
        return
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    # Candidates: local fonts/ dir in repo, /tmp download, system
    pairs = {
        "DejaVuSans": [
            os.path.join(os.path.dirname(__file__), "fonts", "DejaVuSans.ttf"),
            f"{FONT_DIR}/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        ],
        "DejaVuSans-Bold": [
            os.path.join(os.path.dirname(__file__), "fonts", "DejaVuSans-Bold.ttf"),
            f"{FONT_DIR}/DejaVuSans-Bold.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        ],
    }
    for name, paths in pairs.items():
        for p in paths:
            if os.path.exists(p):
                try:
                    pdfmetrics.registerFont(TTFont(name, p))
                    break
                except Exception:
                    continue
    # If Bold not available, alias it to regular so PDF doesn't crash
    try:
        from reportlab.pdfbase.pdfmetrics import getFont
        getFont("DejaVuSans-Bold")
    except Exception:
        try:
            from reportlab.pdfbase.pdfmetrics import getFont
            getFont("DejaVuSans")
            pdfmetrics.registerFontFamily("DejaVuSans",
                normal="DejaVuSans", bold="DejaVuSans",
                italic="DejaVuSans", boldItalic="DejaVuSans")
        except Exception:
            pass
    _fonts_registered = True

SYS = """Ты — эксперт по маркетинговым исследованиям. Отвечай на русском языке.
Для российских рынков используй рубли (руб. млрд), для остальных — доллары ($B).

ТРЕБОВАНИЯ К ДАННЫМ:
- Используй актуальные данные на дату запроса. Не придумывай — если свежих данных нет, укажи последние проверенные с годом.
- Приоритет источников: национальные официальные (Росстат, ЦБ РФ, отраслевые ассоциации для России; Eurostat, нац. статведомства для Европы; Census Bureau, BLS для США).
- Второй приоритет: авторитетные международные (McKinsey, Gartner, IDC, Statista, Bloomberg).
- В конце каждого раздела указывай список источников: [1] Название, год, ссылка."""

# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------
HTML = r"""<!DOCTYPE html>
<html lang="ru">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Market Research</title>
<style>
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;background:#f4f6f9;color:#1a1a1a;min-height:100vh}
.container{max-width:920px;margin:0 auto;padding:28px 16px}
.header{display:flex;align-items:center;gap:12px;margin-bottom:28px}
.logo{width:40px;height:40px;border-radius:10px;background:linear-gradient(135deg,#1D9E75,#0d7a5a);display:flex;align-items:center;justify-content:center;font-size:13px;font-weight:800;color:#fff;flex-shrink:0}
.header-text h1{font-size:20px;font-weight:700;color:#111}
.header-text p{font-size:12px;color:#888;margin-top:2px}
.card{background:#fff;border:1px solid #e4e8ee;border-radius:14px;padding:22px;margin-bottom:16px;box-shadow:0 1px 4px rgba(0,0,0,.04)}
.card-title{font-size:11px;font-weight:700;color:#888;text-transform:uppercase;letter-spacing:.7px;margin-bottom:14px}
.form-grid{display:grid;gap:12px;margin-bottom:12px}
.fg-1{grid-template-columns:1fr}
.fg-3{grid-template-columns:2fr 1fr 1fr}
.fg-2{grid-template-columns:1fr 1fr}
.form-group{display:flex;flex-direction:column;gap:5px}
label{font-size:11px;font-weight:600;color:#666;letter-spacing:.3px}
input,select{padding:9px 11px;font-size:13px;border:1.5px solid #e0e4ea;border-radius:8px;background:#fafbfc;color:#111;transition:border-color .15s,box-shadow .15s;outline:none;width:100%}
input:focus,select:focus{border-color:#1D9E75;box-shadow:0 0 0 3px rgba(29,158,117,.12)}
input::placeholder{color:#c0c4cc}
.tam-wrap{display:flex;gap:6px}
.tam-wrap input{flex:1}
.tam-wrap select{width:100px;flex-shrink:0}
.btn-primary{width:100%;padding:12px;background:linear-gradient(135deg,#1D9E75,#16866a);color:#fff;border:none;border-radius:9px;font-size:14px;font-weight:700;cursor:pointer;transition:opacity .15s;margin-top:4px}
.btn-primary:hover{opacity:.9}
.btn-primary:disabled{opacity:.45;cursor:not-allowed}

/* Progress */
.steps{display:flex;align-items:flex-start;padding:4px 0 8px}
.step{display:flex;flex-direction:column;align-items:center;gap:5px;flex:1}
.step-dot{width:26px;height:26px;border-radius:50%;border:2px solid #dce0e8;background:#f4f6f9;display:flex;align-items:center;justify-content:center;font-size:10px;font-weight:700;color:#aaa;transition:all .3s}
.step-dot.active{border-color:#1D9E75;background:#e6f7f1;color:#0a6b4d}
.step-dot.done{border-color:#1D9E75;background:#1D9E75;color:#fff}
.step-lbl{font-size:10px;color:#bbb;white-space:nowrap;text-align:center;transition:color .3s}
.step-lbl.done{color:#1D9E75;font-weight:600}
.step-line{flex:1;height:2px;background:#e4e8ee;margin-top:-14px;transition:background .4s}
.step-line.done{background:#1D9E75}
@keyframes spin{to{transform:rotate(360deg)}}
.spinner{width:11px;height:11px;border:2px solid rgba(29,158,117,.3);border-top-color:#1D9E75;border-radius:50%;animation:spin .7s linear infinite}
.status{padding:10px 14px;border-radius:9px;font-size:13px;display:flex;align-items:center;gap:8px;margin-top:4px}
.status.loading{background:#e6f7f1;color:#0a6b4d}
.status.error{background:#fdecea;color:#b91c1c}

/* Results */
.result-block{background:#fff;border:1px solid #e4e8ee;border-radius:12px;margin-bottom:12px;overflow:hidden;box-shadow:0 1px 3px rgba(0,0,0,.04)}
.result-head{display:flex;align-items:center;gap:10px;padding:11px 15px;background:#f8fafc;border-bottom:1px solid #e4e8ee;cursor:pointer;user-select:none}
.result-num{width:24px;height:24px;border-radius:50%;background:linear-gradient(135deg,#1D9E75,#16866a);color:#fff;font-size:10px;font-weight:700;display:flex;align-items:center;justify-content:center;flex-shrink:0}
.result-title{font-size:13px;font-weight:600;flex:1;color:#222}
.result-toggle{font-size:11px;color:#bbb}
.result-body{padding:16px;font-size:13px;line-height:1.8;color:#333}
.result-body.hidden{display:none}

/* Structured data inside result */
.data-table{width:100%;border-collapse:collapse;margin:14px 0;font-size:12.5px;box-shadow:0 1px 4px rgba(0,0,0,.06);border-radius:8px;overflow:hidden}
.data-table th{background:#1D9E75;color:#fff;padding:10px 12px;text-align:left;font-weight:600;white-space:nowrap;font-size:12px}
.data-table td{padding:9px 12px;border-bottom:1px solid #eef0f4;vertical-align:top;line-height:1.5}
.data-table tr:last-child td{border-bottom:none}
.data-table tr:nth-child(even) td{background:#f8fafc}
.data-table tr:hover td{background:#f0fdf8}
.share-bar{height:8px;border-radius:4px;background:#1D9E75;opacity:.8;min-width:4px;display:inline-block;vertical-align:middle;margin-right:6px}
.freq-badge{display:inline-block;padding:2px 10px;border-radius:10px;font-size:11px;font-weight:700}
.freq-High{background:#fef3c7;color:#92400e}
.freq-Medium{background:#e0f2fe;color:#075985}
.freq-Low{background:#f0fdf4;color:#166534}
.chart-wrap{margin:14px 0;border-radius:10px;overflow:hidden;border:1px solid #e4e8ee}
.chart-wrap img{width:100%;display:block}
.plain-text{white-space:pre-wrap;word-break:break-word;margin-top:12px;padding-top:12px;border-top:1px solid #f0f0f0;color:#333;font-size:13px;line-height:1.85}
.excel-link{display:inline-flex;align-items:center;gap:6px;margin-top:10px;padding:7px 14px;background:#f0fdf4;border:1px solid #1D9E75;border-radius:8px;color:#0a6b4d;font-size:12px;font-weight:600;cursor:pointer;text-decoration:none}
.excel-link:hover{background:#1D9E75;color:#fff}
.btn-new-report{flex:1;padding:11px 14px;border:none;border-radius:9px;background:linear-gradient(135deg,#f97316,#ea6b0a);color:#fff;font-size:13px;font-weight:700;cursor:pointer;transition:opacity .15s;display:flex;align-items:center;justify-content:center;gap:7px}
.btn-new-report:hover{opacity:.9}

/* Export */
.export-bar{display:flex;gap:10px;margin:4px 0 20px}
.btn-export{flex:1;padding:11px 14px;border:1.5px solid #1D9E75;border-radius:9px;background:#fff;color:#1D9E75;font-size:13px;font-weight:600;cursor:pointer;transition:all .15s;display:flex;align-items:center;justify-content:center;gap:7px}
.btn-export:hover{background:#1D9E75;color:#fff}
.btn-export:disabled{opacity:.4;cursor:not-allowed}

/* Modal */
.modal-overlay{position:fixed;inset:0;background:rgba(0,0,0,.45);display:flex;align-items:center;justify-content:center;z-index:1000;backdrop-filter:blur(3px)}
.modal-overlay.hidden{display:none}
.modal{background:#fff;border-radius:16px;padding:28px;max-width:500px;width:calc(100% - 32px);box-shadow:0 20px 60px rgba(0,0,0,.2)}
.modal-step{font-size:11px;color:#1D9E75;font-weight:700;text-transform:uppercase;letter-spacing:.6px;margin-bottom:8px}
.modal-q{font-size:16px;font-weight:700;color:#111;line-height:1.4;margin-bottom:6px}
.modal-hint{font-size:12px;color:#888;margin-bottom:18px}
.modal-options{display:flex;flex-direction:column;gap:8px;margin-bottom:20px}
.opt-btn{padding:10px 14px;border:1.5px solid #e0e4ea;border-radius:9px;background:#fafbfc;color:#333;font-size:13px;text-align:left;cursor:pointer;transition:all .15s;display:flex;align-items:center;gap:8px}
.opt-btn:hover,.opt-btn.selected{border-color:#1D9E75;background:#f0fdf8;color:#0a6b4d}
.opt-btn.selected{font-weight:600}
.opt-check{width:16px;height:16px;border-radius:4px;border:2px solid #ddd;flex-shrink:0;display:flex;align-items:center;justify-content:center;font-size:10px;transition:all .15s}
.opt-btn.selected .opt-check{background:#1D9E75;border-color:#1D9E75;color:#fff}
.modal-text-input{width:100%;padding:10px 12px;font-size:13px;border:1.5px solid #e0e4ea;border-radius:8px;outline:none;margin-bottom:20px}
.modal-text-input:focus{border-color:#1D9E75;box-shadow:0 0 0 3px rgba(29,158,117,.1)}
.modal-footer{display:flex;gap:10px}
.modal-skip{flex:1;padding:10px;border:1.5px solid #e0e4ea;border-radius:9px;background:#fff;color:#888;font-size:13px;font-weight:600;cursor:pointer}
.modal-skip:hover{border-color:#bbb;color:#555}
.modal-next{flex:2;padding:10px;background:linear-gradient(135deg,#1D9E75,#16866a);color:#fff;border:none;border-radius:9px;font-size:13px;font-weight:700;cursor:pointer}
.modal-next:hover{opacity:.9}
.modal-next:disabled{opacity:.4;cursor:not-allowed}
.modal-loading{display:flex;align-items:center;gap:10px;padding:20px 0;color:#888;font-size:13px}

@media(max-width:600px){.fg-3,.fg-2{grid-template-columns:1fr}}
</style>
</head>
<body>
<div class="container">
  <div class="header">
    <div class="logo">MR</div>
    <div class="header-text">
      <h1>Market Research</h1>
      <p>Создано Treyner Igor &nbsp;·&nbsp; i.treyner@yandex.ru</p>
    </div>
  </div>

  <!-- Form -->
  <div class="card">
    <div class="card-title">Параметры исследования</div>
    <div class="form-grid fg-1">
      <div class="form-group">
        <label>Рынок / тема исследования</label>
        <input id="market" type="text" placeholder="например: рынок электромобилей в России">
      </div>
    </div>
    <div class="form-grid fg-2">
      <div class="form-group">
        <label>Год прогноза</label>
        <input id="year" type="number" value="2030" min="2025" max="2045">
      </div>
      <div class="form-group">
        <label>Тип рынка</label>
        <select id="b2x">
          <option value="B2B и B2C">B2B и B2C</option>
          <option value="B2B">B2B</option>
          <option value="B2C">B2C</option>
        </select>
      </div>
    </div>
    <div class="form-grid fg-2" style="margin-top:4px">
      <div class="form-group">
        <label>География</label>
        <select id="geo">
          <option value="Россия">Россия</option>
          <option value="Global" selected>Global</option>
          <option value="США">США</option>
          <option value="Европа">Европа</option>
          <option value="Азия">Азия</option>
          <option value="СНГ">СНГ</option>
        </select>
      </div>
      <div class="form-group">
        <label>Сегмент / уточнение (необязательно)</label>
        <input id="segment" type="text" placeholder="например: Enterprise SaaS, D2C...">
      </div>
    </div>
    <button class="btn-primary" id="btnStart" onclick="startFlow()">Далее →</button>
  </div>

  <!-- Progress -->
  <div class="card" id="progressCard" style="display:none">
    <div class="steps">
      <div class="step"><div class="step-dot" id="d0">1</div><div class="step-lbl" id="l0">Объём</div></div>
      <div class="step-line" id="ln0"></div>
      <div class="step"><div class="step-dot" id="d1">2</div><div class="step-lbl" id="l1">Игроки</div></div>
      <div class="step-line" id="ln1"></div>
      <div class="step"><div class="step-dot" id="d2">3</div><div class="step-lbl" id="l2">Прогноз</div></div>
      <div class="step-line" id="ln2"></div>
      <div class="step"><div class="step-dot" id="d3">4</div><div class="step-lbl" id="l3">Проблемы</div></div>
      <div class="step-line" id="ln3"></div>
      <div class="step"><div class="step-dot" id="d4">5</div><div class="step-lbl" id="l4">Отчёт</div></div>
    </div>
    <div id="statusBox" class="status loading" style="display:none"></div>
  </div>

  <!-- Export -->
  <div id="exportBar" style="display:none">
    <div class="export-bar">
      <button class="btn-export" id="btnPdf" onclick="exportFile('pdf')">
        <svg width="15" height="15" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="9" y1="13" x2="15" y2="13"/><line x1="9" y1="17" x2="15" y2="17"/></svg>
        Скачать PDF
      </button>
      <button class="btn-export" id="btnWord" onclick="exportFile('word')">
        <svg width="15" height="15" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="8" y1="13" x2="16" y2="13"/><line x1="8" y1="17" x2="16" y2="17"/></svg>
        Скачать Word
      </button>
      <button class="btn-export" id="btnExcel" onclick="exportFile('excel')">
        <svg width="15" height="15" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><rect x="3" y="3" width="18" height="18" rx="2"/><path d="M3 9h18M9 21V9"/></svg>
        Скачать Excel
      </button>
      <button class="btn-new-report" onclick="location.reload()">
        <svg width="15" height="15" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><polyline points="1 4 1 10 7 10"/><path d="M3.51 15a9 9 0 1 0 .49-4.95"/></svg>
        Новый отчёт
      </button>
    </div>
  </div>

  <div id="results"></div>
</div>

<!-- Modal -->
<div class="modal-overlay hidden" id="modalOverlay">
  <div class="modal">
    <div class="modal-step" id="modalStep"></div>
    <div class="modal-q" id="modalQ"></div>
    <div class="modal-hint" id="modalHint"></div>
    <div id="modalBody"></div>
    <div class="modal-footer">
      <button class="modal-skip" onclick="skipQuestion()">Пропустить</button>
      <button class="modal-next" id="modalNext" onclick="nextQuestion()">Далее →</button>
    </div>
  </div>
</div>

<script>
// ── State ──────────────────────────────────────────────────────────────────
const TITLES = ["Объём рынка","Топ игроки","Прогноз роста","Проблемы отрасли","Итоговый отчёт"];
let running = false;
let collectedResults = [];
let structuredData   = {};   // { competitors, growth, problems }
let researchMeta     = {};
let clarifications   = [];   // [{q, a}]
let questionQueue    = [];   // parsed questions from AI
let currentQIdx      = 0;
let currentSelections= [];

// ── Helpers ────────────────────────────────────────────────────────────────
function setStep(i, state) {
  const dot = document.getElementById('d'+i), lbl = document.getElementById('l'+i);
  dot.className = 'step-dot ' + state;
  lbl.className = 'step-lbl ' + (state==='done'?'done':'');
  if (state==='active') dot.innerHTML='<div class="spinner"></div>';
  else if (state==='done'){dot.textContent='✓';if(i<4)document.getElementById('ln'+i).className='step-line done';}
  else dot.textContent=i+1;
}
function showStatus(msg,isErr){
  const el=document.getElementById('statusBox');
  el.style.display='flex';el.className='status '+(isErr?'error':'loading');
  el.innerHTML=isErr?'⚠ '+msg:'<div class="spinner"></div><span>'+msg+'</span>';
}

// ── Modal ──────────────────────────────────────────────────────────────────
function showModal(qObj, idx, total) {
  currentSelections = [];
  document.getElementById('modalOverlay').classList.remove('hidden');
  document.getElementById('modalStep').textContent = `Вопрос ${idx+1} из ${total}`;
  document.getElementById('modalQ').textContent = qObj.question;
  document.getElementById('modalHint').textContent = qObj.hint || '';
  document.getElementById('modalNext').disabled = qObj.type !== 'text';

  const body = document.getElementById('modalBody');
  body.innerHTML = '';

  if (qObj.type === 'single' || qObj.type === 'multi') {
    const wrap = document.createElement('div');
    wrap.className = 'modal-options';
    (qObj.options || []).forEach(opt => {
      const btn = document.createElement('button');
      btn.className = 'opt-btn';
      btn.innerHTML = `<span class="opt-check"></span><span>${opt}</span>`;
      btn.onclick = () => {
        if (qObj.type === 'single') {
          wrap.querySelectorAll('.opt-btn').forEach(b=>b.classList.remove('selected'));
          currentSelections = [opt];
        } else {
          if (currentSelections.includes(opt)) {
            currentSelections = currentSelections.filter(x=>x!==opt);
            btn.classList.remove('selected');
            btn.querySelector('.opt-check').textContent = '';
          } else {
            currentSelections.push(opt);
          }
        }
        btn.classList.toggle('selected', currentSelections.includes(opt));
        if (currentSelections.includes(opt)) btn.querySelector('.opt-check').textContent = '✓';
        document.getElementById('modalNext').disabled = currentSelections.length === 0;
      };
      wrap.appendChild(btn);
    });
    body.appendChild(wrap);
  } else {
    const inp = document.createElement('input');
    inp.type = 'text'; inp.className = 'modal-text-input';
    inp.placeholder = qObj.placeholder || 'Ваш ответ...';
    inp.oninput = () => {
      currentSelections = [inp.value.trim()];
      document.getElementById('modalNext').disabled = !inp.value.trim();
    };
    inp.onkeydown = e => { if(e.key==='Enter' && inp.value.trim()) nextQuestion(); };
    body.appendChild(inp);
    setTimeout(()=>inp.focus(), 100);
  }
}

function hideModal() {
  document.getElementById('modalOverlay').classList.add('hidden');
}

function skipQuestion() {
  clarifications.push({ q: questionQueue[currentQIdx]?.question || '', a: 'пропущено' });
  currentQIdx++;
  proceedModal();
}

function nextQuestion() {
  if (currentSelections.length === 0) return;
  const qObj = questionQueue[currentQIdx];
  clarifications.push({ q: qObj.question, a: currentSelections.join(', ') });
  currentQIdx++;
  proceedModal();
}

function proceedModal() {
  if (currentQIdx >= questionQueue.length) {
    hideModal();
    runResearch();
  } else {
    showModal(questionQueue[currentQIdx], currentQIdx, questionQueue.length);
  }
}

// ── Start flow ─────────────────────────────────────────────────────────────
async function startFlow() {
  const market = document.getElementById('market').value.trim();
  if (!market) { alert('Укажи рынок для исследования'); return; }

  const year    = document.getElementById('year').value || '2030';
  const geo     = document.getElementById('geo').value;
  const b2x     = document.getElementById('b2x').value;
  const segment = document.getElementById('segment').value.trim();
  const cur     = /Россия|СНГ/.test(geo) ? 'руб. млрд' : '$B';
  const segLabel= segment ? `, сегмент: ${segment}` : '';
  const context = `${market} (${geo}${segLabel}, ${b2x})`;

  researchMeta = { market, year, geo, b2x, segment, cur, context };
  clarifications = [];
  currentQIdx = 0;
  questionQueue = [];

  document.getElementById('btnStart').disabled = true;
  document.getElementById('btnStart').textContent = '⏳ Генерирую вопросы...';

  try {
    const resp = await fetch('/clarify', {
      method: 'POST', headers: {'Content-Type':'application/json'},
      body: JSON.stringify({ context })
    });
    const data = await resp.json();
    if (data.error) throw new Error(data.error);
    questionQueue = data.questions || [];
  } catch(e) {
    console.warn('Clarify failed, skip:', e);
    questionQueue = [];
  }

  document.getElementById('btnStart').disabled = false;
  document.getElementById('btnStart').textContent = 'Далее →';

  if (questionQueue.length > 0) {
    showModal(questionQueue[0], 0, questionQueue.length);
  } else {
    runResearch();
  }
}

// ── Research ───────────────────────────────────────────────────────────────
async function runResearch() {
  if (running) return;
  running = true;
  collectedResults = [];
  structuredData   = {};

  document.getElementById('btnStart').disabled = true;
  document.getElementById('results').innerHTML = '';
  document.getElementById('exportBar').style.display = 'none';
  document.getElementById('progressCard').style.display = 'block';
  for (let i=0;i<5;i++) setStep(i,'');

  const { market, year, geo, b2x, cur, context } = researchMeta;
  const clarCtx = clarifications
    .filter(c => c.a !== 'пропущено')
    .map(c => `- ${c.q}: ${c.a}`)
    .join('\n');
  const extra = clarCtx ? `\nДополнительный контекст от пользователя:\n${clarCtx}` : '';

  const prompts = [
    // Step 1 — Market size (plain text ok)
    `Проведи анализ объёма рынка: ${context}.${extra}
Определи TAM и SAM рынка с источниками и годом данных.
## Текущий объём рынка
TAM, SAM, диапазон оценок. Валюта: ${cur}.
## Confidence Level
High / Medium / Low — обоснуй.
## Ключевые источники
3-5 авторитетных источников с годом.`,

    // Step 2 — Competitors → JSON required
    `Топ-10 компаний рынка: ${context}.${extra}
Для каждой: место, компания, доля рынка %, сегмент, регион, ключевое преимущество.
Данные 2023-2025. В конце — 2-3 предложения инсайта.

ОБЯЗАТЕЛЬНО в конце ответа добавь JSON-блок строго в этом формате:
\`\`\`json
{"type":"competitors","data":[{"rank":1,"company":"Name","share":25.0,"segment":"...","region":"...","note":"..."}]}
\`\`\``,

    // Step 3 — Growth → JSON required
    `Прогноз роста рынка: ${context} до ${year}.${extra}
Три сценария с CAGR и итогом в ${cur}:
- Пессимистичный: CAGR X%, итог Y ${cur} — причины.
- Базовый: CAGR X%, итог Y ${cur}.
- Оптимистичный: CAGR X%, итог Y ${cur}.
Топ-5 драйверов роста и топ-3 риска.

ОБЯЗАТЕЛЬНО в конце добавь JSON-блок строго в этом формате:
\`\`\`json
{"type":"growth","baseYear":2024,"targetYear":${year},"baseSize":0,"currency":"${cur}","scenarios":{"pessimistic":{"cagr":5,"finalSize":0},"base":{"cagr":10,"finalSize":0},"optimistic":{"cagr":18,"finalSize":0}},"drivers":["драйвер1","драйвер2","драйвер3"]}
\`\`\`
Заполни finalSize реальными числами из расчёта.`,

    // Step 4 — Problems → JSON required
    `Ключевые проблемы рынка: ${context}.${extra}
Минимум 8 проблем по категориям: Технологические, Регуляторные, Рыночные, Операционные, Финансовые.
Для каждой: описание, частота (High/Medium/Low), источники.
В конце: топ-3 возможности для новых игроков.

ОБЯЗАТЕЛЬНО в конце добавь JSON-блок строго в этом формате:
\`\`\`json
{"type":"problems","data":[{"category":"Рыночные","title":"...","description":"...","frequency":"High","sources":["..."]}]}
\`\`\``,

    // Step 5 — Final report
    `Итоговый стратегический отчёт: ${context}, прогноз до ${year}.${extra}
## TL;DR
3-4 ключевых вывода одной фразой каждый.
## Обзор рынка
## Размер рынка (${cur})
## Конкурентный ландшафт
## Прогноз роста
## Ключевые проблемы
## Стратегический вывод
ИДТИ / НЕ ИДТИ — с чётким обоснованием и рекомендуемой точкой входа.`
  ];

  for (let i=0; i<prompts.length; i++) {
    setStep(i,'active');
    showStatus(`Шаг ${i+1}/5 — ${TITLES[i]}...`);
    try {
      const resp = await fetch('/research', {
        method:'POST', headers:{'Content-Type':'application/json'},
        body: JSON.stringify({ prompt: prompts[i] })
      });
      const data = await resp.json();
      if (data.error) throw new Error(data.error);
      setStep(i,'done');
      const parsed = parseStructured(data.result);
      if (parsed) Object.assign(structuredData, { [parsed.type]: parsed });
      renderResult(i, data.result, parsed);
      collectedResults.push({ title: TITLES[i], content: data.result, structured: parsed });
    } catch(e) {
      setStep(i,'');
      showStatus('Ошибка на шаге '+(i+1)+': '+e.message, true);
      running=false; document.getElementById('btnStart').disabled=false; return;
    }
  }

  document.getElementById('statusBox').style.display='none';
  document.getElementById('exportBar').style.display='block';
  running=false;
  document.getElementById('btnStart').disabled=false;
  document.getElementById('btnStart').textContent='Новое исследование →';
}

// ── Parse JSON from response ───────────────────────────────────────────────
function parseStructured(text) {
  const m = text.match(/```json\n([\s\S]*?)\n```/);
  if (!m) return null;
  try { return JSON.parse(m[1]); } catch(e) { return null; }
}

function stripJson(text) {
  return text.replace(/```json\n[\s\S]*?\n```/g, '').trim();
}

// ── Render result block ────────────────────────────────────────────────────
function renderResult(i, rawText, structured) {
  const div = document.createElement('div');
  div.className = 'result-block';

  const plainText = stripJson(rawText);
  let innerHtml = `<div class="plain-text">${escHtml(plainText)}</div>`;

  if (structured) {
    if (structured.type === 'competitors') {
      innerHtml = renderCompetitorsTable(structured.data) + innerHtml;
    } else if (structured.type === 'growth') {
      innerHtml = renderGrowthChart(structured) + innerHtml;
    } else if (structured.type === 'problems') {
      innerHtml = renderProblemsTable(structured.data) + innerHtml;
    }
  }

  div.innerHTML = `
    <div class="result-head" onclick="tog(${i})">
      <div class="result-num">${i+1}</div>
      <div class="result-title">${TITLES[i]}</div>
      <div class="result-toggle" id="t${i}">▲</div>
    </div>
    <div class="result-body" id="b${i}">${innerHtml}</div>`;
  document.getElementById('results').appendChild(div);
}

function tog(i) {
  const b=document.getElementById('b'+i), t=document.getElementById('t'+i);
  b.classList.toggle('hidden'); t.textContent=b.classList.contains('hidden')?'▼':'▲';
}

function escHtml(t) {
  return t.replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');
}

// Competitors table
function renderCompetitorsTable(data) {
  if (!data || !data.length) return '';
  const maxShare = Math.max(...data.map(r=>parseFloat(r.share)||0));
  let rows = data.map(r => {
    const share = parseFloat(r.share)||0;
    const barW  = maxShare > 0 ? Math.round((share/maxShare)*80) : 0;
    const pct   = share > 0 ? share+'%' : (r.share||'—');
    return `<tr>
      <td style="font-weight:700;color:#1D9E75;text-align:center;width:36px">${r.rank}</td>
      <td style="font-weight:700;color:#111;min-width:140px">${escHtml(r.company||'')}</td>
      <td style="min-width:110px"><div style="display:flex;align-items:center;gap:6px">
        <div style="width:${barW}px;height:8px;background:#1D9E75;border-radius:4px;opacity:.8;flex-shrink:0;min-width:4px"></div>
        <span style="font-weight:600;white-space:nowrap">${pct}</span></div></td>
      <td style="color:#444">${escHtml(r.segment||'')}</td>
      <td style="color:#444;white-space:nowrap">${escHtml(r.region||'')}</td>
      <td style="color:#555;font-size:12px;line-height:1.55">${escHtml(r.note||'')}</td>
    </tr>`;
  }).join('');
  return `<div>
    <div style="overflow-x:auto">
      <table class="data-table" style="min-width:680px">
        <thead><tr><th>#</th><th>Компания</th><th>Доля рынка</th><th>Сегмент</th><th>Регион</th><th>Ключевое преимущество</th></tr></thead>
        <tbody>${rows}</tbody>
      </table>
    </div>
    <button class="excel-link" onclick="exportFile('excel')">
      <svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><rect x="3" y="3" width="18" height="18" rx="2"/><path d="M3 9h18M9 21V9"/></svg>
      Скачать таблицу (Полная версия в Excel)
    </button>
  </div>`;
}

// Problems table
function renderProblemsTable(data) {
  if (!data || !data.length) return '';
  let rows = data.map(p => `<tr>
    <td style="font-weight:700;color:#1D9E75;white-space:nowrap;min-width:100px">${escHtml(p.category||'')}</td>
    <td style="font-weight:600;color:#111;min-width:120px">${escHtml(p.title||'')}</td>
    <td style="color:#444;font-size:13px;line-height:1.6">${escHtml(p.description||'')}</td>
    <td style="text-align:center;white-space:nowrap"><span class="freq-badge freq-${p.frequency}">${p.frequency}</span></td>
  </tr>`).join('');
  return `<div>
    <div style="overflow-x:auto">
      <table class="data-table" style="min-width:600px">
        <thead><tr><th>Категория</th><th>Проблема</th><th>Описание</th><th>Частота</th></tr></thead>
        <tbody>${rows}</tbody>
      </table>
    </div>
    <button class="excel-link" onclick="exportFile('excel')">
      <svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><rect x="3" y="3" width="18" height="18" rx="2"/><path d="M3 9h18M9 21V9"/></svg>
      Скачать таблицу (Полная версия в Excel)
    </button>
  </div>`;
}

// Growth chart using Canvas
function renderGrowthChart(g) {
  const id = 'chart_' + Date.now();
  // Draw after DOM insert
  setTimeout(() => drawGrowthChart(id, g), 50);
  return `<div class="chart-wrap"><canvas id="${id}" height="260" style="width:100%"></canvas></div>`;
}

function drawGrowthChart(id, g) {
  const canvas = document.getElementById(id);
  if (!canvas) return;
  canvas.width = canvas.offsetWidth || 800;
  const ctx = canvas.getContext('2d');
  const W = canvas.width, H = canvas.height;
  const pad = { top:30, right:30, bottom:50, left:60 };
  const cw = W - pad.left - pad.right, ch = H - pad.top - pad.bottom;

  const baseYear = g.baseYear || 2024;
  const targetYear = g.targetYear || 2030;
  const baseSize = g.baseSize || 50;
  const sc = g.scenarios || {};

  const years = [];
  for (let y = baseYear; y <= targetYear; y++) years.push(y);

  function project(cagr) {
    return years.map((_, i) => +(baseSize * Math.pow(1 + cagr/100, i)).toFixed(1));
  }

  const pess = project(sc.pessimistic?.cagr || 0);
  const base = project(sc.base?.cagr || 0);
  const opti = project(sc.optimistic?.cagr || 0);

  const allVals = [...pess, ...base, ...opti];
  const minV = 0, maxV = Math.max(...allVals) * 1.1;

  function xPos(i) { return pad.left + (i / (years.length-1)) * cw; }
  function yPos(v) { return pad.top + ch - ((v - minV) / (maxV - minV)) * ch; }

  // Background
  ctx.fillStyle = '#fff';
  ctx.fillRect(0, 0, W, H);

  // Grid lines
  ctx.strokeStyle = '#f0f0f0'; ctx.lineWidth = 1;
  for (let t=0; t<=5; t++) {
    const y = pad.top + (t/5)*ch;
    ctx.beginPath(); ctx.moveTo(pad.left, y); ctx.lineTo(pad.left+cw, y); ctx.stroke();
    const val = maxV - (t/5)*(maxV-minV);
    ctx.fillStyle='#aaa'; ctx.font='10px sans-serif'; ctx.textAlign='right';
    ctx.fillText(val.toFixed(0), pad.left-6, y+4);
  }

  // Currency label
  ctx.fillStyle='#888'; ctx.font='10px sans-serif'; ctx.textAlign='left';
  ctx.fillText(g.currency||'', 4, pad.top-8);

  // Year labels
  years.forEach((yr, i) => {
    ctx.fillStyle='#888'; ctx.font='10px sans-serif'; ctx.textAlign='center';
    ctx.fillText(yr, xPos(i), H-12);
  });

  // Area fill for base scenario
  ctx.beginPath();
  ctx.moveTo(xPos(0), yPos(pess[0]));
  pess.forEach((v,i) => ctx.lineTo(xPos(i), yPos(v)));
  opti.slice().reverse().forEach((v,i) => ctx.lineTo(xPos(opti.length-1-i), yPos(v)));
  ctx.closePath();
  ctx.fillStyle = 'rgba(29,158,117,0.08)';
  ctx.fill();

  // Draw lines
  function drawLine(vals, color, dash=[]) {
    ctx.beginPath(); ctx.strokeStyle=color; ctx.lineWidth=2.5;
    ctx.setLineDash(dash);
    vals.forEach((v,i) => i===0 ? ctx.moveTo(xPos(i),yPos(v)) : ctx.lineTo(xPos(i),yPos(v)));
    ctx.stroke(); ctx.setLineDash([]);
    // Dot at end
    ctx.beginPath(); ctx.arc(xPos(vals.length-1), yPos(vals[vals.length-1]), 4, 0, Math.PI*2);
    ctx.fillStyle=color; ctx.fill();
  }

  drawLine(pess, '#f59e0b', [5,4]);
  drawLine(opti, '#60a5fa', [5,4]);
  drawLine(base, '#1D9E75');

  // Legend
  const legend = [
    { label:`Базовый (CAGR ${sc.base?.cagr||0}%)`, color:'#1D9E75', dash:false },
    { label:`Оптимистичный (${sc.optimistic?.cagr||0}%)`, color:'#60a5fa', dash:true },
    { label:`Пессимистичный (${sc.pessimistic?.cagr||0}%)`, color:'#f59e0b', dash:true },
  ];
  legend.forEach((l,i) => {
    const lx = pad.left + i*(cw/3), ly = pad.top - 16;
    ctx.strokeStyle=l.color; ctx.lineWidth=2.5;
    if(l.dash) ctx.setLineDash([5,4]); else ctx.setLineDash([]);
    ctx.beginPath(); ctx.moveTo(lx,ly); ctx.lineTo(lx+22,ly); ctx.stroke();
    ctx.setLineDash([]);
    ctx.fillStyle='#444'; ctx.font='10px sans-serif'; ctx.textAlign='left';
    ctx.fillText(l.label, lx+26, ly+4);
  });
}

// ── Export ─────────────────────────────────────────────────────────────────
async function exportFile(type) {
  const btnId = type==='pdf'?'btnPdf': type==='word'?'btnWord':'btnExcel';
  const btn = document.getElementById(btnId);
  const orig = btn ? btn.innerHTML : '';
  if(btn){ btn.disabled=true; btn.innerHTML='<span>⏳ Генерирую...</span>'; }

  // Capture chart as base64 if exists
  const chartCanvas = document.querySelector('canvas[id^="chart_"]');
  const chartImg = chartCanvas ? chartCanvas.toDataURL('image/png') : null;

  try {
    const resp = await fetch('/export-'+type, {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({ results:collectedResults, meta:researchMeta, chartImg, clarifications })
    });
    if (!resp.ok) { const e=await resp.json().catch(()=>({})); throw new Error(e.error||'Ошибка'); }
    const blob = await resp.blob();
    const url  = URL.createObjectURL(blob);
    const a    = document.createElement('a');
    a.href=url;
    const safe=(researchMeta.market||'research').replace(/[^a-zA-Z\u0400-\u04FF0-9]/g,'_').slice(0,40);
    a.download=`market_research_${safe}.${type==='pdf'?'pdf':'docx'}`;
    a.click(); URL.revokeObjectURL(url);
  } catch(e) { alert('Ошибка экспорта: '+e.message); }
  finally { if(btn){ btn.disabled=false; btn.innerHTML=orig; } }
}
</script>
</body>
</html>"""


# ---------------------------------------------------------------------------
# Perplexity helper
# ---------------------------------------------------------------------------
async def _perplexity(prompt: str, sys_override: str = None) -> str:
    async with httpx.AsyncClient(timeout=90) as client:
        resp = await client.post(
            PERPLEXITY_URL,
            headers={"Authorization": f"Bearer {PERPLEXITY_API_KEY}", "Content-Type": "application/json"},
            json={
                "model": "sonar-pro",
                "messages": [
                    {"role": "system", "content": sys_override or SYS},
                    {"role": "user",   "content": prompt},
                ],
                "max_tokens": 2000,
                "temperature": 0.2,
            },
        )
        resp.raise_for_status()
        return resp.json()["choices"][0]["message"]["content"]


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------
@app.route("/")
def index():
    return HTML


@app.route("/research", methods=["POST"])
def research():
    prompt = request.json.get("prompt", "")
    try:
        result = asyncio.run(_perplexity(prompt))
        return jsonify({"result": result})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/clarify", methods=["POST"])
def clarify():
    """Generate up to 3 clarifying questions based on the market context."""
    context = request.json.get("context", "")
    sys_clarify = """Ты — помощник по маркетинговым исследованиям.
Твоя задача: на основе описания рынка сгенерировать до 3 уточняющих вопросов,
которые помогут сделать исследование точнее.

Отвечай ТОЛЬКО валидным JSON без лишнего текста:
{
  "questions": [
    {
      "question": "Текст вопроса",
      "hint": "Краткое пояснение зачем этот вопрос",
      "type": "single",
      "options": ["Вариант 1", "Вариант 2", "Вариант 3"]
    }
  ]
}

Типы вопросов:
- "single" — один вариант из списка (дай 3-5 опций)
- "multi" — несколько вариантов (дай 4-6 опций)
- "text" — свободный ответ (без поля options)

Вопросы должны быть конкретными и релевантными именно этому рынку.
Максимум 3 вопроса. Если рынок очень понятный — можно 1-2.
Примеры тем: целевой сегмент клиентов, стадия развития продукта,
главный конкурентный фокус, приоритет географии.

ВАЖНО: если задаёшь вопрос о регионе/географии — первым вариантом ВСЕГДА ставь
"Все регионы / вся страна", затем конкретные регионы."""

    prompt = f"Рынок для исследования: {context}\n\nСгенерируй уточняющие вопросы."
    try:
        raw = asyncio.run(_perplexity(prompt, sys_clarify))
        # Strip markdown fences if present
        raw = re.sub(r"```json\s*|\s*```", "", raw).strip()
        data = json.loads(raw)
        questions = data.get("questions", [])[:3]
        return jsonify({"questions": questions})
    except Exception as e:
        return jsonify({"questions": [], "error": str(e)})


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------

@app.route("/export-excel", methods=["POST"])
def export_excel():
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        payload = request.json
        results = payload.get("results", [])
        meta    = payload.get("meta", {})
        market  = meta.get("market", "Исследование рынка")
        geo     = meta.get("geo", "")
        year    = meta.get("year", "")

        wb = openpyxl.Workbook()
        wb.remove(wb.active)

        GREEN_FILL = PatternFill("solid", fgColor="1D9E75")
        ALT_FILL   = PatternFill("solid", fgColor="F4F6F9")
        H_FONT     = Font(name="Calibri", bold=True, color="FFFFFF", size=11)
        B_FONT     = Font(name="Calibri", bold=True, size=10)
        N_FONT     = Font(name="Calibri", size=10)
        BORDER     = Border(
            left=Side(style="thin", color="D0D5DD"),
            right=Side(style="thin", color="D0D5DD"),
            top=Side(style="thin", color="D0D5DD"),
            bottom=Side(style="thin", color="D0D5DD"),
        )

        def style_header(ws, cols):
            for c in range(1, cols+1):
                cell = ws.cell(row=1, column=c)
                cell.font = H_FONT
                cell.fill = GREEN_FILL
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                cell.border = BORDER
            ws.row_dimensions[1].height = 28

        def style_row(ws, row_idx, cols, alt=False):
            for c in range(1, cols+1):
                cell = ws.cell(row=row_idx, column=c)
                cell.font = N_FONT
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border = BORDER
                if alt:
                    cell.fill = ALT_FILL

        # ── Summary sheet
        ws0 = wb.create_sheet("Сводка")
        ws0.column_dimensions["A"].width = 22
        ws0.column_dimensions["B"].width = 40
        ws0["A1"] = "Market Research Report"
        ws0["A1"].font = Font(name="Calibri", bold=True, size=14, color="1D9E75")
        ws0.merge_cells("A1:B1")
        for i, (k, v) in enumerate([("Рынок", market), ("География", geo),
                                      ("Прогноз до", year),
                                      ("Дата", datetime.now().strftime("%d.%m.%Y"))], 2):
            ws0.cell(row=i, column=1, value=k).font = B_FONT
            ws0.cell(row=i, column=2, value=v).font = N_FONT

        # ── Find structured data in results
        for section in results:
            structured = section.get("structured")
            if not structured:
                continue
            stype = structured.get("type")

            if stype == "competitors" and structured.get("data"):
                ws = wb.create_sheet("Топ игроки")
                headers = ["#", "Компания", "Доля рынка %", "Сегмент", "Регион", "Ключевое преимущество"]
                for ci, h in enumerate(headers, 1):
                    ws.cell(row=1, column=ci, value=h)
                style_header(ws, len(headers))
                for ri, r in enumerate(structured["data"], 2):
                    vals = [r.get("rank",""), r.get("company",""), r.get("share",""),
                            r.get("segment",""), r.get("region",""), r.get("note","")]
                    for ci, v in enumerate(vals, 1):
                        ws.cell(row=ri, column=ci, value=v)
                    style_row(ws, ri, len(headers), alt=(ri%2==0))
                ws.column_dimensions["A"].width = 5
                ws.column_dimensions["B"].width = 28
                ws.column_dimensions["C"].width = 12
                ws.column_dimensions["D"].width = 20
                ws.column_dimensions["E"].width = 18
                ws.column_dimensions["F"].width = 45
                ws.freeze_panes = "A2"

            elif stype == "growth":
                ws = wb.create_sheet("Прогноз роста")
                sc        = structured.get("scenarios", {})
                base_size = structured.get("baseSize", 0)
                base_year = structured.get("baseYear", 2024)
                tgt_year  = structured.get("targetYear", int(year) if year else 2030)
                cur       = structured.get("currency", "$B")
                headers   = ["Год", f"Пессимист. ({cur})", f"Базовый ({cur})", f"Оптимист. ({cur})"]
                for ci, h in enumerate(headers, 1):
                    ws.cell(row=1, column=ci, value=h)
                style_header(ws, 4)
                for i, yr in enumerate(range(base_year, int(tgt_year)+1)):
                    ri = i + 2
                    p = round(base_size*(1+sc.get("pessimistic",{}).get("cagr",0)/100)**i, 1)
                    b = round(base_size*(1+sc.get("base",{}).get("cagr",0)/100)**i, 1)
                    o = round(base_size*(1+sc.get("optimistic",{}).get("cagr",0)/100)**i, 1)
                    for ci, v in enumerate([yr, p, b, o], 1):
                        ws.cell(row=ri, column=ci, value=v)
                    style_row(ws, ri, 4, alt=(i%2==0))
                # Scenario summary
                ws.append([])
                ws.append(["Сценарий", "CAGR", f"Итог {tgt_year}"])
                for key, label in [("pessimistic","Пессимист."),("base","Базовый"),("optimistic","Оптимист.")]:
                    s = sc.get(key, {})
                    ws.append([label, f"{s.get('cagr','')}%", s.get("finalSize","")])
                for col in ["A","B","C","D"]:
                    ws.column_dimensions[col].width = 18
                ws.freeze_panes = "A2"

            elif stype == "problems" and structured.get("data"):
                ws = wb.create_sheet("Проблемы отрасли")
                headers = ["Категория", "Проблема", "Описание", "Частота", "Источники"]
                for ci, h in enumerate(headers, 1):
                    ws.cell(row=1, column=ci, value=h)
                style_header(ws, len(headers))
                for ri, p in enumerate(structured["data"], 2):
                    vals = [p.get("category",""), p.get("title",""), p.get("description",""),
                            p.get("frequency",""), ", ".join(p.get("sources",[]))]
                    for ci, v in enumerate(vals, 1):
                        ws.cell(row=ri, column=ci, value=v)
                    style_row(ws, ri, len(headers), alt=(ri%2==0))
                ws.column_dimensions["A"].width = 18
                ws.column_dimensions["B"].width = 22
                ws.column_dimensions["C"].width = 52
                ws.column_dimensions["D"].width = 10
                ws.column_dimensions["E"].width = 30
                ws.freeze_panes = "A2"

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        safe_name = re.sub(r"[^\w]", "_", market)[:40]
        return send_file(buf,
                         mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                         as_attachment=True,
                         download_name=f"market_research_{safe_name}.xlsx")
    except Exception as e:
        import traceback
        return jsonify({"error": str(e), "trace": traceback.format_exc()}), 500

@app.route("/export-pdf", methods=["POST"])
def export_pdf():
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                        HRFlowable, Table, TableStyle, Image)
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

        register_pdf_fonts()
        NRM  = "DejaVuSans"
        BLD  = "DejaVuSans-Bold"

        payload        = request.json
        results        = payload.get("results", [])
        meta           = payload.get("meta", {})
        clarifications = payload.get("clarifications", [])
        chart_img_b64  = payload.get("chartImg")

        market = meta.get("market", "Исследование рынка")
        geo    = meta.get("geo", "")
        year   = meta.get("year", "")
        b2x    = meta.get("b2x", "")
        tamCur = "руб. млрд" if meta.get("geo","") in ("Россия","СНГ") else "$B"

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2.2*cm, bottomMargin=2*cm,
                                title=f"Market Research: {market}")

        GREEN  = colors.HexColor("#1D9E75")
        DARK   = colors.HexColor("#111111")
        GRAY   = colors.HexColor("#666666")
        LGRAY  = colors.HexColor("#f4f6f9")
        HGRAY  = colors.HexColor("#e8ecf0")
        BORDER = colors.HexColor("#d0d5dd")
        WHITE  = colors.white

        def ps(name, **kw): return ParagraphStyle(name, **kw)
        st_title  = ps("T",  fontSize=22, leading=28, textColor=DARK,  fontName=BLD, spaceAfter=4)
        st_market = ps("Mk", fontSize=16, leading=22, textColor=GREEN, fontName=BLD, spaceAfter=12)
        st_meta   = ps("Me", fontSize=10, leading=14, textColor=GRAY,  fontName=NRM, spaceAfter=3)
        st_h1     = ps("H1", fontSize=14, leading=20, textColor=GREEN, fontName=BLD, spaceBefore=18, spaceAfter=6)
        st_h2     = ps("H2", fontSize=12, leading=16, textColor=DARK,  fontName=BLD, spaceBefore=10, spaceAfter=4)
        st_h3     = ps("H3", fontSize=11, leading=15, textColor=DARK,  fontName=BLD, spaceBefore=6,  spaceAfter=3)
        st_body   = ps("B",  fontSize=10, leading=15, textColor=colors.HexColor("#333"),
                        fontName=NRM, spaceAfter=3, alignment=TA_JUSTIFY)
        st_bullet = ps("Bu", fontSize=10, leading=15, textColor=colors.HexColor("#333"),
                        fontName=NRM, spaceAfter=2, leftIndent=14)
        st_foot   = ps("F",  fontSize=8,  leading=12, textColor=GRAY, fontName=NRM, alignment=TA_CENTER)
        st_clabel = ps("CL", fontSize=9,  leading=13, textColor=GRAY, fontName=NRM, spaceAfter=2)
        st_th     = ps("TH", fontSize=9,  leading=12, textColor=WHITE, fontName=BLD)
        st_td     = ps("TD", fontSize=9,  leading=12, textColor=DARK,  fontName=NRM)

        def cell(text, style):
            return Paragraph(str(text), style)

        BASE_TBL_STYLE = [
            ("BACKGROUND",   (0,0), (-1,0),  GREEN),
            ("TEXTCOLOR",    (0,0), (-1,0),  WHITE),
            ("FONTNAME",     (0,0), (-1,0),  BLD),
            ("FONTSIZE",     (0,0), (-1,0),  9),
            ("FONTNAME",     (0,1), (-1,-1), NRM),
            ("FONTSIZE",     (0,1), (-1,-1), 9),
            ("ROWBACKGROUNDS",(0,1),(-1,-1), [WHITE, LGRAY]),
            ("GRID",         (0,0), (-1,-1), 0.5, BORDER),
            ("TOPPADDING",   (0,0), (-1,-1), 5),
            ("BOTTOMPADDING",(0,0), (-1,-1), 5),
            ("LEFTPADDING",  (0,0), (-1,-1), 7),
            ("RIGHTPADDING", (0,0), (-1,-1), 7),
            ("VALIGN",       (0,0), (-1,-1), "TOP"),
            ("LINEBELOW",    (0,0), (-1,0),  1.5, GREEN),
        ]

        story = []

        # Cover
        story.append(Paragraph("Market Research", st_title))
        story.append(Paragraph(market, st_market))
        meta_items = []
        if geo:  meta_items.append(f"География: {geo}")
        if b2x:  meta_items.append(f"Тип рынка: {b2x}")
        if year: meta_items.append(f"Прогноз до: {year}")
        meta_items.append(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
        for m in meta_items:
            story.append(Paragraph(m, st_meta))

        real_cl = [c for c in clarifications if c.get("a") != "пропущено"]
        if real_cl:
            story.append(Spacer(1, 6))
            story.append(Paragraph("Уточнения:", st_meta))
            for c in real_cl:
                story.append(Paragraph(f"• {c['q']}: {c['a']}", st_clabel))

        story.append(Spacer(1, 12))
        story.append(HRFlowable(width="100%", thickness=2, color=GREEN, spaceAfter=20))

        # Sections
        for idx, section in enumerate(results):
            title      = section.get("title", f"Раздел {idx+1}")
            content    = section.get("content", "")
            structured = section.get("structured")

            story.append(Paragraph(f"{idx+1}. {title}", st_h1))
            story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER, spaceAfter=8))

            if structured:
                stype = structured.get("type")

                if stype == "competitors" and structured.get("data"):
                    rows = structured["data"]
                    tbl_data = [[cell(h, st_th) for h in ["#","Компания","Доля","Сегмент","Регион","Преимущество"]]]
                    for r in rows:
                        tbl_data.append([
                            cell(str(r.get("rank","")),    st_td),
                            cell(r.get("company",""),      st_td),
                            cell(f"{r.get('share','')}%",  st_td),
                            cell(r.get("segment",""),      st_td),
                            cell(r.get("region",""),       st_td),
                            cell(r.get("note",""),         st_td),
                        ])
                    tbl = Table(tbl_data, colWidths=[0.7*cm, 3.5*cm, 1.5*cm, 2.8*cm, 2.3*cm, 4.7*cm],
                                repeatRows=1)
                    tbl.setStyle(TableStyle(BASE_TBL_STYLE))
                    story.append(tbl)
                    story.append(Spacer(1, 10))

                elif stype == "growth":
                    sc  = structured.get("scenarios", {})
                    ty  = structured.get("targetYear", year)
                    cur = structured.get("currency", tamCur)
                    tbl_data = [[cell(h, st_th) for h in ["Сценарий", "CAGR", f"Итог {ty} ({cur})"]]]
                    for label, key in [("Пессимистичный","pessimistic"),("Базовый","base"),("Оптимистичный","optimistic")]:
                        s = sc.get(key, {})
                        tbl_data.append([cell(label, st_td), cell(f"{s.get('cagr','')}%", st_td),
                                         cell(str(s.get("finalSize","")), st_td)])
                    tbl = Table(tbl_data, colWidths=[5*cm, 3*cm, 7.5*cm], repeatRows=1)
                    tbl.setStyle(TableStyle(BASE_TBL_STYLE))
                    story.append(tbl)
                    story.append(Spacer(1, 8))

                    if chart_img_b64 and idx == 2:
                        try:
                            img_data = base64.b64decode(chart_img_b64.split(",")[-1])
                            story.append(Image(io.BytesIO(img_data), width=15*cm, height=6*cm))
                            story.append(Spacer(1, 8))
                        except Exception:
                            pass

                    drivers = structured.get("drivers", [])
                    if drivers:
                        story.append(Paragraph("Драйверы роста: " + " · ".join(drivers), st_body))
                    story.append(Spacer(1, 6))

                elif stype == "problems" and structured.get("data"):
                    rows = structured["data"]
                    tbl_data = [[cell(h, st_th) for h in ["Категория","Проблема","Описание","Частота"]]]
                    for p in rows:
                        tbl_data.append([
                            cell(p.get("category",""),    st_td),
                            cell(p.get("title",""),       st_td),
                            cell(p.get("description",""), st_td),
                            cell(p.get("frequency",""),   st_td),
                        ])
                    tbl = Table(tbl_data, colWidths=[2.8*cm, 3.2*cm, 7.5*cm, 2*cm], repeatRows=1)
                    tbl.setStyle(TableStyle(BASE_TBL_STYLE))
                    story.append(tbl)
                    story.append(Spacer(1, 10))


            # Plain text
            plain = re.sub(r"```json[\s\S]*?```", "", content).strip()
            for line in plain.split("\n"):
                line = line.rstrip()
                if not line:
                    story.append(Spacer(1, 4)); continue
                clean = re.sub(r"\*\*|\*", "", line)
                if line.startswith("## "):
                    story.append(Paragraph(clean[3:], st_h2))
                elif line.startswith("### "):
                    story.append(Paragraph(clean[4:], st_h3))
                elif line.startswith("- ") or line.startswith("* "):
                    story.append(Paragraph("• " + clean[2:], st_bullet))
                else:
                    story.append(Paragraph(clean, st_body))
            story.append(Spacer(1, 8))

        story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER, spaceBefore=14, spaceAfter=6))
        story.append(Paragraph(
            f"Treyner Igor · i.treyner@yandex.ru · {datetime.now().strftime('%d.%m.%Y %H:%M')}",
            st_foot))

        doc.build(story)
        buf.seek(0)
        safe_name = re.sub(r"[^\w]", "_", market)[:40]
        return send_file(buf, mimetype="application/pdf", as_attachment=True,
                         download_name=f"market_research_{safe_name}.pdf")
    except Exception as e:
        import traceback
        return jsonify({"error": str(e), "trace": traceback.format_exc()}), 500


@app.route("/export-word", methods=["POST"])
def export_word():
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        payload        = request.json
        results        = payload.get("results", [])
        meta           = payload.get("meta", {})
        clarifications = payload.get("clarifications", [])
        chart_img_b64  = payload.get("chartImg")

        market = meta.get("market", "Исследование рынка")
        geo    = meta.get("geo", "")
        year   = meta.get("year", "")
        b2x    = meta.get("b2x", "")
        tamCur = "руб. млрд" if meta.get("geo","") in ("Россия","СНГ") else "$B"

        doc = DocxDocument()
        for sec in doc.sections:
            sec.top_margin = sec.bottom_margin = Cm(2.5)
            sec.left_margin = sec.right_margin = Cm(2.5)

        GREEN = RGBColor(0x1D, 0x9E, 0x75)
        DARK  = RGBColor(0x11, 0x11, 0x11)
        GRAY  = RGBColor(0x66, 0x66, 0x66)

        def hr(color_hex="1D9E75", sz=12):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), str(sz))
            bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), color_hex)
            pBdr.append(bot); pPr.append(pBdr)
            p.paragraph_format.space_after = Pt(4)

        def add_cell_text(cell, text, bold=False, size=10, color=None):
            for para in cell.paragraphs: para.clear()
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.bold = bold; r.font.size = Pt(size)
            if color: r.font.color.rgb = color

        # Title
        p = doc.add_paragraph()
        r = p.add_run("Market Research\n"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=DARK
        r2 = p.add_run(market); r2.bold=True; r2.font.size=Pt(16); r2.font.color.rgb=GREEN
        doc.add_paragraph()

        # Meta table
        rows = []
        if geo:  rows.append(("География", geo))
        if b2x:  rows.append(("Тип рынка", b2x))
        if year: rows.append(("Прогноз до", year))
        rows.append(("Дата", datetime.now().strftime("%d.%m.%Y")))
        real_cl = [c for c in clarifications if c.get('a') != 'пропущено']
        for c in real_cl:
            rows.append((c['q'][:40], c['a']))

        tbl = doc.add_table(rows=len(rows), cols=2)
        tbl.style = "Table Grid"
        for i, (k, v) in enumerate(rows):
            tbl.rows[i].cells[0].width = Cm(4)
            tbl.rows[i].cells[1].width = Cm(12)
            add_cell_text(tbl.rows[i].cells[0], k, bold=True, color=GRAY)
            add_cell_text(tbl.rows[i].cells[1], v)

        doc.add_paragraph()
        hr("1D9E75", 14)
        doc.add_paragraph()

        # Sections
        for idx, section in enumerate(results):
            title      = section.get("title", f"Раздел {idx+1}")
            content    = section.get("content", "")
            structured = section.get("structured")

            h = doc.add_heading(f"{idx+1}. {title}", level=1)
            for run in h.runs:
                run.font.color.rgb = GREEN; run.font.size = Pt(14)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after  = Pt(6)
            hr("e0e4ea", 4)

            if structured:
                stype = structured.get("type")

                if stype == "competitors" and structured.get("data"):
                    data_rows = structured["data"]
                    headers = ["#", "Компания", "Доля", "Сегмент", "Регион", "Преимущество"]
                    t = doc.add_table(rows=len(data_rows)+1, cols=6)
                    t.style = "Table Grid"
                    for ci, h_txt in enumerate(headers):
                        add_cell_text(t.rows[0].cells[ci], h_txt, bold=True, size=9, color=RGBColor(0x1D,0x9E,0x75))
                    for ri, r_data in enumerate(data_rows):
                        vals = [str(r_data.get("rank","")), r_data.get("company",""),
                                f"{r_data.get('share','')}%", r_data.get("segment",""),
                                r_data.get("region",""), r_data.get("note","")]
                        for ci, v in enumerate(vals):
                            add_cell_text(t.rows[ri+1].cells[ci], v, size=9)
                    doc.add_paragraph()

                elif stype == "growth":
                    sc  = structured.get("scenarios", {})
                    ty  = structured.get("targetYear", year)
                    cur = structured.get("currency", tamCur)
                    headers2 = ["Сценарий", "CAGR", f"Итог {ty} ({cur})"]
                    sc_rows  = [
                        ["Пессимистичный", f"{sc.get('pessimistic',{}).get('cagr','')}%", str(sc.get('pessimistic',{}).get('finalSize',''))],
                        ["Базовый",        f"{sc.get('base',{}).get('cagr','')}%",        str(sc.get('base',{}).get('finalSize',''))],
                        ["Оптимистичный",  f"{sc.get('optimistic',{}).get('cagr','')}%",  str(sc.get('optimistic',{}).get('finalSize',''))],
                    ]
                    t2 = doc.add_table(rows=4, cols=3)
                    t2.style = "Table Grid"
                    for ci, h_txt in enumerate(headers2):
                        add_cell_text(t2.rows[0].cells[ci], h_txt, bold=True, size=10, color=RGBColor(0x1D,0x9E,0x75))
                    for ri, row_vals in enumerate(sc_rows):
                        for ci, v in enumerate(row_vals):
                            add_cell_text(t2.rows[ri+1].cells[ci], v, size=10)
                    doc.add_paragraph()

                    # Chart image
                    if chart_img_b64 and idx == 2:
                        try:
                            img_data = base64.b64decode(chart_img_b64.split(",")[-1])
                            img_buf  = io.BytesIO(img_data)
                            doc.add_picture(img_buf, width=Cm(14))
                        except Exception:
                            pass
                    doc.add_paragraph()

                elif stype == "problems" and structured.get("data"):
                    data_rows = structured["data"]
                    headers3  = ["Категория", "Проблема", "Описание", "Частота"]
                    t3 = doc.add_table(rows=len(data_rows)+1, cols=4)
                    t3.style = "Table Grid"
                    for ci, h_txt in enumerate(headers3):
                        add_cell_text(t3.rows[0].cells[ci], h_txt, bold=True, size=9, color=RGBColor(0x1D,0x9E,0x75))
                    for ri, p_data in enumerate(data_rows):
                        vals = [p_data.get("category",""), p_data.get("title",""),
                                p_data.get("description",""), p_data.get("frequency","")]
                        for ci, v in enumerate(vals):
                            add_cell_text(t3.rows[ri+1].cells[ci], v, size=9)
                    doc.add_paragraph()

            # Plain text
            plain = re.sub(r"```json\n[\s\S]*?\n```", "", content).strip()
            for line in plain.split("\n"):
                line = line.rstrip()
                if not line:
                    doc.add_paragraph().paragraph_format.space_after = Pt(2); continue
                if line.startswith("## "):
                    h2 = doc.add_heading(re.sub(r"\*\*|\*","",line[3:]), level=2)
                    for run in h2.runs: run.font.size=Pt(12); run.font.color.rgb=DARK
                    h2.paragraph_format.space_before=Pt(8)
                elif line.startswith("### "):
                    h3 = doc.add_heading(re.sub(r"\*\*|\*","",line[4:]), level=3)
                    for run in h3.runs: run.font.size=Pt(11)
                    h3.paragraph_format.space_before=Pt(6)
                elif line.startswith("- ") or line.startswith("* "):
                    p = doc.add_paragraph(style="List Bullet")
                    for part in re.split(r"(\*\*.*?\*\*)", line[2:]):
                        rn = p.add_run(part[2:-2] if part.startswith("**") and part.endswith("**") else part)
                        rn.bold = part.startswith("**"); rn.font.size = Pt(10)
                    p.paragraph_format.space_after = Pt(2)
                else:
                    p = doc.add_paragraph()
                    for part in re.split(r"(\*\*.*?\*\*)", line):
                        rn = p.add_run(part[2:-2] if part.startswith("**") and part.endswith("**") else part)
                        rn.bold = part.startswith("**"); rn.font.size = Pt(10)
                    p.paragraph_format.space_after = Pt(2)
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

        hr("e0e4ea", 4)
        fp = doc.add_paragraph()
        r = fp.add_run(f"Treyner Igor · i.treyner@yandex.ru · {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        r.font.size=Pt(8); r.font.color.rgb=GRAY
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buf = io.BytesIO()
        doc.save(buf); buf.seek(0)
        safe_name = re.sub(r"[^\w]", "_", market)[:40]
        return send_file(buf,
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         as_attachment=True,
                         download_name=f"market_research_{safe_name}.docx")
    except Exception as e:
        import traceback
        return jsonify({"error": str(e), "trace": traceback.format_exc()}), 500


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    print(f"\n  Market Research Server → http://localhost:{port}\n")
    app.run(debug=False, host="0.0.0.0", port=port)
